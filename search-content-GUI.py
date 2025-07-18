import os
import json
import re
import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox
from pathlib import Path
import threading
from datetime import datetime


class SearchContentGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Content Search Tool")
        self.root.geometry("800x600")
        self.root.configure(bg="#f0f0f0")

        # Variables
        self.search_term = tk.StringVar()
        self.case_sensitive = tk.BooleanVar(value=False)
        self.is_searching = False

        self.setup_ui()

    def setup_ui(self):
        # Title
        title_label = tk.Label(
            self.root,
            text="Content Search Tool",
            font=("Arial", 16, "bold"),
            bg="#f0f0f0",
        )
        title_label.pack(pady=10)

        # Search Frame
        search_frame = tk.Frame(self.root, bg="#f0f0f0")
        search_frame.pack(fill="x", padx=20, pady=10)

        # Search term input
        tk.Label(search_frame, text="Search Term:", bg="#f0f0f0").pack(anchor="w")
        self.search_entry = tk.Entry(
            search_frame, textvariable=self.search_term, font=("Arial", 12), width=60
        )
        self.search_entry.pack(fill="x", pady=5)

        # Options frame
        options_frame = tk.Frame(search_frame, bg="#f0f0f0")
        options_frame.pack(fill="x", pady=5)

        # Case sensitive checkbox
        tk.Checkbutton(
            options_frame,
            text="Case Sensitive",
            variable=self.case_sensitive,
            bg="#f0f0f0",
        ).pack(side="left")

        # Search button
        self.search_button = tk.Button(
            options_frame,
            text="Search",
            command=self.start_search,
            bg="#4CAF50",
            fg="white",
            font=("Arial", 10, "bold"),
        )
        self.search_button.pack(side="right", padx=5)

        # Progress bar
        self.progress = ttk.Progressbar(search_frame, mode="indeterminate")
        self.progress.pack(fill="x", pady=5)

        # Main content frame
        content_frame = tk.Frame(self.root, bg="#f0f0f0")
        content_frame.pack(fill="both", expand=True, padx=20, pady=10)

        # Logs frame
        logs_frame = tk.LabelFrame(
            content_frame, text="Search Logs", font=("Arial", 10, "bold"), bg="#f0f0f0"
        )
        logs_frame.pack(fill="both", expand=True, pady=(0, 10))

        # Logs text area
        self.logs_text = scrolledtext.ScrolledText(
            logs_frame, height=8, font=("Courier", 9)
        )
        self.logs_text.pack(fill="both", expand=True, padx=5, pady=5)

        # Results frame
        results_frame = tk.LabelFrame(
            content_frame,
            text="Search Results",
            font=("Arial", 10, "bold"),
            bg="#f0f0f0",
        )
        results_frame.pack(fill="both", expand=True)

        # Results text area
        self.results_text = scrolledtext.ScrolledText(
            results_frame, height=8, font=("Courier", 9)
        )
        self.results_text.pack(fill="both", expand=True, padx=5, pady=5)

        # Bind Enter key to search
        self.search_entry.bind("<Return>", lambda event: self.start_search())

        # Focus on search entry
        self.search_entry.focus()

    def log_message(self, message):
        """Add message to logs area"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_entry = f"[{timestamp}] {message}\n"
        self.logs_text.insert(tk.END, log_entry)
        self.logs_text.see(tk.END)
        self.root.update_idletasks()

    def clear_logs(self):
        """Clear logs area"""
        self.logs_text.delete(1.0, tk.END)

    def clear_results(self):
        """Clear results area"""
        self.results_text.delete(1.0, tk.END)

    def start_search(self):
        """Start the search in a separate thread"""
        if self.is_searching:
            return

        search_term = self.search_term.get().strip()
        if not search_term:
            messagebox.showwarning("Input Error", "Please enter a search term.")
            return

        # Clear previous results
        self.clear_logs()
        self.clear_results()

        # Start search in separate thread
        self.is_searching = True
        self.search_button.config(state="disabled", text="Searching...")
        self.progress.start()

        search_thread = threading.Thread(
            target=self.perform_search, args=(search_term, self.case_sensitive.get())
        )
        search_thread.daemon = True
        search_thread.start()

    def perform_search(self, search_term, case_sensitive):
        """Perform the actual search"""
        try:
            self.log_message(f"Starting search for: '{search_term}'")
            self.log_message(f"Case sensitive: {case_sensitive}")
            self.log_message("=" * 50)

            matching_files = self.search_files_for_term(
                ".", search_term, case_sensitive
            )

            # Update results on main thread
            self.root.after(0, self.display_results, matching_files, search_term)

        except Exception as e:
            self.root.after(0, self.search_error, str(e))

    def display_results(self, matching_files, search_term):
        """Display search results"""
        self.progress.stop()
        self.search_button.config(state="normal", text="Search")
        self.is_searching = False

        if not matching_files:
            self.log_message("Search completed - No files found")
            self.results_text.insert(
                tk.END, "No files found containing the search term.\n"
            )
            return

        self.log_message(f"Search completed - Found {len(matching_files)} file(s)")

        # Display results
        self.results_text.insert(
            tk.END, f"Found {len(matching_files)} file(s) containing '{search_term}':\n"
        )
        self.results_text.insert(tk.END, "-" * 60 + "\n")

        for file_path in matching_files:
            file_name = os.path.basename(file_path)
            directory = os.path.dirname(file_path)
            result_line = f"File: {file_name} | Directory: {directory}\n"
            self.results_text.insert(tk.END, result_line)

    def search_error(self, error_message):
        """Handle search errors"""
        self.progress.stop()
        self.search_button.config(state="normal", text="Search")
        self.is_searching = False

        self.log_message(f"Search error: {error_message}")
        messagebox.showerror(
            "Search Error", f"An error occurred during search:\n{error_message}"
        )

    # Your existing search functions (modified to work with GUI)
    def search_in_text_file(self, file_path, search_term, case_sensitive=False):
        """Search for a term in plain text files (.py, .md)"""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                content = file.read()
                if case_sensitive:
                    return search_term in content
                else:
                    return search_term.lower() in content.lower()
        except Exception as e:
            self.log_message(f"Error reading {file_path}: {e}")
            return False

    def search_in_jupyter_notebook(self, file_path, search_term, case_sensitive=False):
        """Search for a term in Jupyter notebook files (.ipynb)"""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                notebook = json.load(file)

            # Extract text from all cells
            text_content = []
            for cell in notebook.get("cells", []):
                if cell.get("cell_type") == "code":
                    source = cell.get("source", [])
                    if isinstance(source, list):
                        text_content.extend(source)
                    else:
                        text_content.append(source)
                elif cell.get("cell_type") == "markdown":
                    source = cell.get("source", [])
                    if isinstance(source, list):
                        text_content.extend(source)
                    else:
                        text_content.append(source)

            full_text = "".join(text_content)

            if case_sensitive:
                return search_term in full_text
            else:
                return search_term.lower() in full_text.lower()

        except Exception as e:
            self.log_message(f"Error reading Jupyter notebook {file_path}: {e}")
            return False

    def search_in_word_doc(self, file_path, search_term, case_sensitive=False):
        """Search for a term in Word documents (.doc, .docx)"""
        try:
            from docx import Document

            doc = Document(file_path)
            full_text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)

            content = "\n".join(full_text)

            if case_sensitive:
                return search_term in content
            else:
                return search_term.lower() in content.lower()

        except ImportError:
            self.log_message(
                "python-docx not installed. Install with: pip install python-docx"
            )
            return False
        except Exception as e:
            self.log_message(f"Error reading Word document {file_path}: {e}")
            return False

    def search_in_pdf(self, file_path, search_term, case_sensitive=False):
        """Search for a term in PDF files using pdfplumber (.pdf)"""
        try:
            import pdfplumber

            full_text = []

            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:  # Only add non-empty text
                        full_text.append(text)

            content = "\n".join(full_text)

            if case_sensitive:
                return search_term in content
            else:
                return search_term.lower() in content.lower()

        except ImportError:
            self.log_message(
                "pdfplumber not installed. Install with: pip install pdfplumber"
            )
            return False
        except Exception as e:
            self.log_message(f"Error reading PDF {file_path}: {e}")
            return False

    def search_files_for_term(self, base_dir, search_term, case_sensitive=False):
        """Search for a term in specified file types within a directory."""
        extensions = {
            ".py": self.search_in_text_file,
            ".md": self.search_in_text_file,
            ".ipynb": self.search_in_jupyter_notebook,
            ".doc": self.search_in_word_doc,
            ".docx": self.search_in_word_doc,
            ".pdf": self.search_in_pdf,
        }

        # Folders to exclude from search
        excluded_folders = {
            "venv99855",
            ".venv99855",
            "env",
            ".env",
            "node_modules",
            ".git",
            "__pycache__",
            ".pytest_cache",
            ".mypy_cache",
            "site-packages",
            ".ipynb_checkpoints",
        }

        matching_files = []

        for root, dirs, files in os.walk(base_dir):
            # Remove excluded directories from dirs list (case-insensitive)
            dirs[:] = [d for d in dirs if d.lower() not in excluded_folders]

            # Log directories being searched
            if dirs:
                self.log_message(f"Searching directories: {dirs}")

            for file in files:
                file_path = os.path.join(root, file)
                file_extension = Path(file).suffix.lower()

                if file_extension in extensions:
                    search_function = extensions[file_extension]
                    if search_function(file_path, search_term, case_sensitive):
                        matching_files.append(file_path)

        return matching_files


def main():
    root = tk.Tk()
    app = SearchContentGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
